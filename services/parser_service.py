# 多模态解析服务
import os
import io
import base64
from typing import Dict, Any, Optional, List
from fastapi import UploadFile, HTTPException
import PyPDF2
import docx
from PIL import Image
import pytesseract
from pdf2image import convert_from_bytes
import requests
from config import settings
from services.embedding_service import embedding_service
from services.database_service import database_service

class ParserService:
    """
    多模态解析服务 - 处理文本、图片、音频、文档等不同格式的文件
    """
    
    def __init__(self):
        self.supported_types = {
            'pdf': self._parse_pdf,
            'docx': self._parse_docx,
            'txt': self._parse_text,
            'md': self._parse_text,
            'png': self._parse_image,
            'jpg': self._parse_image,
            'jpeg': self._parse_image,
            'gif': self._parse_image
        }
        self.max_tokens = 6000  # 设置最大token数，留一些余量
    
    def _split_text_into_chunks(self, text: str, max_tokens: int = None) -> List[str]:
        """将长文本分割成多个块"""
        if max_tokens is None:
            max_tokens = self.max_tokens
        
        # 更保守的token估算：中文字符1.2倍，英文单词1.5倍
        def estimate_tokens(text: str) -> int:
            chinese_chars = len([c for c in text if '\u4e00' <= c <= '\u9fff'])
            english_words = len(text.split())
            # 更保守的估算
            return int(chinese_chars * 1.2 + english_words * 1.5)
        
        if estimate_tokens(text) <= max_tokens:
            return [text]
        
        # 按段落分割
        paragraphs = text.split('\n\n')
        chunks = []
        current_chunk = ""
        
        for paragraph in paragraphs:
            # 如果单个段落就超过限制，需要进一步分割
            if estimate_tokens(paragraph) > max_tokens:
                # 先保存当前块
                if current_chunk:
                    chunks.append(current_chunk.strip())
                    current_chunk = ""
                
                # 分割过长的段落
                sub_chunks = self._split_long_paragraph(paragraph, max_tokens)
                chunks.extend(sub_chunks)
            elif estimate_tokens(current_chunk + paragraph) <= max_tokens:
                current_chunk += paragraph + "\n\n"
            else:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                current_chunk = paragraph + "\n\n"
        
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        return chunks
    
    def _split_long_paragraph(self, paragraph: str, max_tokens: int) -> List[str]:
        """分割过长的段落"""
        def estimate_tokens(text: str) -> int:
            chinese_chars = len([c for c in text if '\u4e00' <= c <= '\u9fff'])
            english_words = len(text.split())
            return int(chinese_chars * 1.2 + english_words * 1.5)
        
        # 按句子分割
        sentences = paragraph.split('。')
        chunks = []
        current_chunk = ""
        
        for sentence in sentences:
            if not sentence.strip():
                continue
                
            if estimate_tokens(current_chunk + sentence) <= max_tokens:
                current_chunk += sentence + "。"
            else:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                current_chunk = sentence + "。"
        
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        return chunks
    
    async def parse_file(self, file: UploadFile, s3_data: Dict[str, Any], user_id: str) -> Dict[str, Any]:
        """
        解析上传的文件并创建记忆单元
        
        Args:
            file: 上传的文件
            s3_data: S3上传返回的数据
            user_id: 用户ID
            
        Returns:
            Dict: 解析结果和记忆ID
        """
        try:
            file_extension = file.filename.split('.')[-1].lower()
            
            if file_extension not in self.supported_types:
                raise HTTPException(
                    status_code=400,
                    detail=f"Unsupported file type: {file_extension}"
                )
            
            # 重置文件指针
            await file.seek(0)
            file_content = await file.read()
            
            # 根据文件类型解析
            parser_func = self.supported_types[file_extension]
            parsed_content = await parser_func(file_content, file)
            
            # 生成embedding
            embedding = embedding_service.generate_embedding(
                text=parsed_content['text'],
                input_type="passage"
            )
            
            # 创建记忆单元
            memory_id = database_service.create_memory(
                content=parsed_content['text'],
                memory_type=parsed_content['type'],
                embedding=embedding,
                user_id=user_id,
                metadata={
                    'original_filename': file.filename,
                    'file_size': len(file_content),
                    'file_extension': file_extension,
                    's3_key': s3_data.get('s3_key'),
                    's3_url': s3_data.get('file_url'),
                    **parsed_content.get('metadata', {})
                },
                source=s3_data.get('file_url'),
                summary=parsed_content.get('summary'),
                tags=parsed_content.get('tags', [])
            )
            # 处理额外的文本块（如果有的话）
            additional_memories = []
            if 'additional_chunks' in parsed_content:
                for i, chunk in enumerate(parsed_content['additional_chunks']):
                    try:
                        # 为每个块生成embedding
                        chunk_embedding = embedding_service.generate_embedding(
                            text=chunk,
                            input_type="passage"
                        )
                        
                        # ✅ 创建额外的记忆单元 - 添加 user_id 参数
                        chunk_memory_id = database_service.create_memory(
                            content=chunk,
                            memory_type=parsed_content['type'],
                            embedding=chunk_embedding,
                            user_id=user_id,  # ✅ 添加 user_id 参数
                            metadata={
                                'original_filename': file.filename,
                                'file_size': len(file_content),
                                'file_extension': file_extension,
                                's3_key': s3_data.get('s3_key'),
                                's3_url': s3_data.get('file_url'),
                                'chunk_index': i + 1,
                                'total_chunks': parsed_content['metadata'].get('total_chunks', 1),
                                'is_partial': True,
                                **parsed_content.get('metadata', {})
                            },
                            source=s3_data.get('file_url'),
                            summary=chunk[:200] + "..." if len(chunk) > 200 else chunk,
                            tags=parsed_content.get('tags', [])
                        )
                        
                        additional_memories.append({
                            'memory_id': chunk_memory_id,
                            'chunk_index': i + 1,
                            'content_preview': chunk[:100] + "..." if len(chunk) > 100 else chunk
                        })
                        
                    except Exception as e:
                        print(f"⚠️  Failed to process chunk {i + 1}: {e}")
                        import traceback
                        traceback.print_exc()
                        continue
            
            return {
                'success': True,
                'memory_id': memory_id,
                'parsed_content': parsed_content,
                'embedding_dimension': len(embedding),
                'additional_memories': additional_memories
            }
            
        except Exception as e:
            print(f"❌ File parsing failed: {e}")
            raise HTTPException(status_code=500, detail=f"Parsing failed: {str(e)}")
    
    async def _parse_pdf(self, content: bytes, file: UploadFile) -> Dict[str, Any]:
        """解析PDF文件"""
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
            text = ""
            
            # 首先尝试直接提取文本
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text()
                text += page_text + "\n"
            
            # 清理文本
            text = text.strip()
            
            # 如果文本为空或很短，尝试OCR
            if not text or len(text) < 50:
                print("📄 PDF文本为空，尝试OCR识别...")
                try:
                    # 将PDF转换为图片
                    images = convert_from_bytes(content)
                    ocr_text = ""
                    
                    for i, image in enumerate(images):
                        print(f"🔍 OCR处理第{i+1}页...")
                        # 使用pytesseract进行OCR
                        page_text = pytesseract.image_to_string(image, lang='chi_sim+eng')
                        ocr_text += page_text + "\n"
                    
                    if ocr_text.strip():
                        text = ocr_text.strip()
                        print(f"✅ OCR成功，提取文本长度: {len(text)}")
                    else:
                        print("⚠️ OCR也未提取到文本")
                        
                except Exception as ocr_error:
                    print(f"❌ OCR失败: {ocr_error}")
                    # OCR失败时，返回提示信息而不是空文本
                    text = f"[扫描版PDF - 需要OCR识别] 此PDF文件是扫描版，无法直接提取文本。请安装OCR依赖或使用可编辑的PDF文件。错误详情: {str(ocr_error)}"
            
            # 检查文本长度，如果太长则分块
            text_chunks = self._split_text_into_chunks(text)
            
            if len(text_chunks) > 1:
                # 如果文本被分块，返回第一个块，其他块作为额外信息
                return {
                    'text': text_chunks[0],
                    'type': 'document',
                    'metadata': {
                        'page_count': len(pdf_reader.pages),
                        'parser': 'PyPDF2+OCR' if len(text) > 50 else 'PyPDF2',
                        'total_chunks': len(text_chunks),
                        'chunk_index': 0,
                        'is_partial': True
                    },
                    'summary': text[:200] + "..." if len(text) > 200 else text,
                    'tags': ['pdf', 'document'],
                    'additional_chunks': text_chunks[1:]  # 其他块
                }
            else:
                return {
                    'text': text,
                    'type': 'document',
                    'metadata': {
                        'page_count': len(pdf_reader.pages),
                        'parser': 'PyPDF2+OCR' if len(text) > 50 else 'PyPDF2'
                    },
                    'summary': text[:200] + "..." if len(text) > 200 else text,
                    'tags': ['pdf', 'document']
                }
            
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"PDF parsing failed: {str(e)}")
    
    async def _parse_docx(self, content: bytes, file: UploadFile) -> Dict[str, Any]:
        """解析DOCX文件"""
        try:
            doc = docx.Document(io.BytesIO(content))
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # 清理文本
            text = text.strip()
            
            return {
                'text': text,
                'type': 'document',
                'metadata': {
                    'paragraph_count': len(doc.paragraphs),
                    'parser': 'python-docx'
                },
                'summary': text[:200] + "..." if len(text) > 200 else text,
                'tags': ['docx', 'document']
            }
            
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"DOCX parsing failed: {str(e)}")
    
    async def _parse_text(self, content: bytes, file: UploadFile) -> Dict[str, Any]:
        """解析文本文件"""
        try:
            # 尝试不同的编码
            encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
            text = ""
            
            for encoding in encodings:
                try:
                    text = content.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue
            
            if not text:
                text = content.decode('utf-8', errors='ignore')
            
            # 清理文本
            text = text.strip()
            
            return {
                'text': text,
                'type': 'text',
                'metadata': {
                    'encoding': 'utf-8',
                    'parser': 'builtin'
                },
                'summary': text[:200] + "..." if len(text) > 200 else text,
                'tags': ['text', 'plain']
            }
            
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Text parsing failed: {str(e)}")
    
    async def _parse_image(self, content: bytes, file: UploadFile) -> Dict[str, Any]:
        """解析图片文件"""
        try:
            # 打开图片
            image = Image.open(io.BytesIO(content))
            
            # 获取图片信息
            width, height = image.size
            format_name = image.format
            mode = image.mode
            
            # 生成图片描述（这里简化处理，实际应该使用CLIP或OCR）
            image_description = f"Image: {width}x{height} pixels, format: {format_name}, mode: {mode}"
            
            # 将图片转换为base64用于存储
            img_buffer = io.BytesIO()
            image.save(img_buffer, format=format_name)
            img_base64 = base64.b64encode(img_buffer.getvalue()).decode()
            
            return {
                'text': image_description,
                'type': 'image',
                'metadata': {
                    'width': width,
                    'height': height,
                    'format': format_name,
                    'mode': mode,
                    'base64_data': img_base64,
                    'parser': 'PIL'
                },
                'summary': f"Image file: {file.filename} ({width}x{height})",
                'tags': ['image', format_name.lower()]
            }
            
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Image parsing failed: {str(e)}")
    
    async def parse_text_input(self, text: str, source: Optional[str], user_id: str) -> Dict[str, Any]:
        """解析纯文本输入并创建记忆"""
        try:
            # 生成embedding
            embedding = embedding_service.generate_embedding(
                text=text,
                input_type="passage"
            )
            
            # ✅ 正确的方式 - 按照Lin的版本
            memory_id = database_service.create_memory(
                content=text,              # 第1个参数
                memory_type='text',        # 第2个参数
                embedding=embedding,       # 第3个参数
                user_id=user_id,          # 第4个参数 ✅
                metadata={                 # 第5个参数（命名参数）
                    'source': source or 'direct_input',
                    'parser': 'direct_text'
                },
                source=source,             # 第6个参数（命名参数）
                summary=text[:200] + "..." if len(text) > 200 else text,  # 第7个参数
                tags=['text', 'direct_input']  # 第8个参数 - 你可以用 [] 或 ['text', 'direct_input']
            )
            
            return {
                'success': True,
                'memory_id': memory_id,
                'content': text,
                'embedding_dimension': len(embedding)
            }
            
        except Exception as e:
            print(f"❌ Text parsing failed: {e}")
            import traceback
            traceback.print_exc()
            raise HTTPException(status_code=500, detail=f"Text parsing failed: {str(e)}")
        
    def get_supported_types(self) -> List[str]:
        """获取支持的文件类型"""
        return list(self.supported_types.keys())
    
    def health_check(self) -> Dict[str, Any]:
        """健康检查"""
        try:
            # 检查依赖库
            import PyPDF2
            import docx
            from PIL import Image
            
            return {
                'status': 'healthy',
                'supported_types': self.get_supported_types(),
                'dependencies': {
                    'PyPDF2': 'available',
                    'python-docx': 'available',
                    'PIL': 'available'
                }
            }
            
        except ImportError as e:
            return {
                'status': 'unhealthy',
                'error': f"Missing dependency: {str(e)}",
                'supported_types': []
            }

# 全局实例
parser_service = ParserService()
